# Copyright 2025 Databend Labs
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import io
import json
from pathlib import Path

from databend_aiserver.udfs.files import _read_docx, _read_pdf
from databend_aiserver.udfs.docparse import ai_parse_document
from databend_aiserver.stages.operator import get_operator, resolve_stage_subpath
from pypdf import PdfReader, PdfWriter
from docx import Document
from docx.enum.text import WD_BREAK

DATA_DIR = Path(__file__).resolve().parents[1] / "data"


def _write_multipage_pdf(stage, filename: str = "multi.pdf"):
    base = (DATA_DIR / "sample.pdf").read_bytes()
    reader = PdfReader(io.BytesIO(base))
    writer = PdfWriter()
    for _ in range(2):
        for page in reader.pages:
            writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    operator = get_operator(stage)
    operator.write(resolve_stage_subpath(stage, filename), buf.getvalue())


def _write_multipage_docx(stage, filename: str = "multi.docx"):
    doc = Document()
    doc.add_heading("Page One", level=1)
    doc.add_paragraph("Content of page one.")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("Page Two", level=1)
    doc.add_paragraph("Content of page two.")
    buf = io.BytesIO()
    doc.save(buf)
    operator = get_operator(stage)
    operator.write(resolve_stage_subpath(stage, filename), buf.getvalue())


def test_read_pdf(memory_stage):
    result = _read_pdf(memory_stage, "sample.pdf")

    assert isinstance(result, str)
    assert "dummy" in result.replace(" ", "").lower()


def test_read_docx(memory_stage):
    result = _read_docx(memory_stage, "sample.docx")

    assert isinstance(result, str)
    assert "This is a short paragraph" in result


def test_parse_document(memory_stage):
    _write_multipage_pdf(memory_stage, "multi.pdf")
    result = ai_parse_document(memory_stage, "multi.pdf")

    assert isinstance(result, dict)
    normalized = dict(result)
    normalized["metadata"] = {"pageCount": "<PAGECOUNT>"}
    normalized["pages"] = [
        {"index": p["index"], "content": " ".join(p["content"].split())}
        for p in (normalized.get("pages") or [])
    ]

    expected = {
        "pages": [
            {"index": 0, "content": "Dumm y PDF file"},
            {"index": 1, "content": "Dumm y PDF file"},
        ],
        "metadata": {"pageCount": "<PAGECOUNT>"},
        "errorInformation": None,
    }
    actual_str = json.dumps(normalized, ensure_ascii=False, sort_keys=True)
    expected_str = json.dumps(expected, ensure_ascii=False, sort_keys=True)
    assert actual_str == expected_str


def test_parse_document_docx(memory_stage):
    _write_multipage_docx(memory_stage, "multi.docx")
    result = ai_parse_document(memory_stage, "multi.docx")

    assert isinstance(result, dict)
    normalized = dict(result)
    normalized["metadata"] = {"pageCount": "<PAGECOUNT>"}
    normalized["pages"] = [
        {"index": p["index"], "content": " ".join(p["content"].split())}
        for p in (normalized.get("pages") or [])
    ]
    expected = {
        "pages": [
            {"index": 0, "content": "Page One Content of page one."},
            {"index": 1, "content": "Page Two Content of page two."},
        ],
        "metadata": {"pageCount": "<PAGECOUNT>"},
        "errorInformation": None,
    }
    actual_str = json.dumps(normalized, ensure_ascii=False, sort_keys=True)
    expected_str = json.dumps(expected, ensure_ascii=False, sort_keys=True)
    assert actual_str == expected_str
