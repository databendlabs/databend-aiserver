# Copyright 2025 Databend Labs
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import io
import json
from pathlib import Path

from databend_udf.client import UDFClient
from docx import Document
from docx.enum.text import WD_BREAK
from pypdf import PdfReader, PdfWriter

from tests.integration.conftest import build_stage_mapping
from databend_aiserver.stages.operator import get_operator, resolve_stage_subpath


DATA_DIR = Path(__file__).resolve().parents[1] / "data"
PDF_SRC = DATA_DIR / "2206.01062.pdf"
DOCX_SRC = DATA_DIR / "lorem_ipsum.docx"


def _write_multipage_pdf(stage, filename: str):
    base = PDF_SRC.read_bytes()
    reader = PdfReader(io.BytesIO(base))
    writer = PdfWriter()
    for _ in range(2):
        for page in reader.pages:
            writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    operator = get_operator(stage)
    operator.write(resolve_stage_subpath(stage, filename), buf.getvalue())


def _write_multipage_docx(stage, filename: str):
    doc = Document(DOCX_SRC)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("Second page - appended for test.")
    buf = io.BytesIO()
    doc.save(buf)
    operator = get_operator(stage)
    operator.write(resolve_stage_subpath(stage, filename), buf.getvalue())


def _call_docparse(client: UDFClient, path: str, memory_stage):
    result = client.call_function(
        "ai_parse_document",
        path,
        stage_locations=[build_stage_mapping(memory_stage)],
    )
    assert len(result) == 1
    payload_raw = result[0]
    if isinstance(payload_raw, (bytes, bytearray)):
        payload_raw = payload_raw.decode("utf-8")
    if isinstance(payload_raw, str):
        payload = json.loads(payload_raw)
    else:
        payload = payload_raw
    return payload


def _normalize_payload(payload):
    pages = [
        {"index": p["index"], "content": "<PAGE_CONTENT>"}
        for p in (payload.get("pages") or [])
    ]
    return {
        "pages": pages,
        "metadata": {"pageCount": "<PAGECOUNT>"},
        "errorInformation": payload.get("errorInformation"),
    }


def test_parse_document_pdf_round_trip(running_server, memory_stage):
    _write_multipage_pdf(memory_stage, "multi.pdf")
    client = UDFClient(host="127.0.0.1", port=running_server)

    payload = _call_docparse(client, "multi.pdf", memory_stage)

    normalized = _normalize_payload(payload)
    page_count = len(normalized["pages"])
    expected = {
        "pages": [{"index": i, "content": "<PAGE_CONTENT>"} for i in range(page_count)],
        "metadata": {"pageCount": "<PAGECOUNT>"},
        "errorInformation": None,
    }

    assert json.dumps(normalized, sort_keys=True) == json.dumps(expected, sort_keys=True)


def test_parse_document_docx_round_trip(running_server, memory_stage):
    _write_multipage_docx(memory_stage, "multi.docx")
    client = UDFClient(host="127.0.0.1", port=running_server)

    payload = _call_docparse(client, "multi.docx", memory_stage)

    normalized = _normalize_payload(payload)
    page_count = len(normalized["pages"])
    expected = {
        "pages": [{"index": i, "content": "<PAGE_CONTENT>"} for i in range(page_count)],
        "metadata": {"pageCount": "<PAGECOUNT>"},
        "errorInformation": None,
    }

    assert json.dumps(normalized, sort_keys=True) == json.dumps(expected, sort_keys=True)
