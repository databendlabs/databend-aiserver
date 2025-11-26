# Copyright 2025 Databend Labs
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import json

from databend_udf.client import UDFClient

from tests.integration.conftest import build_stage_mapping
from databend_aiserver.stages.operator import get_operator, resolve_stage_subpath
from pypdf import PdfReader, PdfWriter
from docx import Document
from docx.enum.text import WD_BREAK
from pathlib import Path
import io

DATA_DIR = Path(__file__).resolve().parents[1] / "data"
PDF_SRC = DATA_DIR / "2206.01062.pdf"
DOCX_SRC = DATA_DIR / "lorem_ipsum.docx"


def _write_multipage_pdf(stage, filename: str):
    base = PDF_SRC.read_bytes()
    reader = PdfReader(io.BytesIO(base))
    writer = PdfWriter()
    for _ in range(2):
        for page in reader.pages:
            writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    operator = get_operator(stage)
    operator.write(resolve_stage_subpath(stage, filename), buf.getvalue())


def _write_multipage_docx(stage, filename: str):
    doc = Document(DOCX_SRC)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("Second page - appended for test.")
    buf = io.BytesIO()
    doc.save(buf)
    operator = get_operator(stage)
    operator.write(resolve_stage_subpath(stage, filename), buf.getvalue())


def test_parse_document_round_trip(running_server, memory_stage):
    _write_multipage_pdf(memory_stage, "multi.pdf")
    client = UDFClient(host="127.0.0.1", port=running_server)
    result = client.call_function(
        "ai_parse_document",
        "multi.pdf",
        2048,
        stage_locations=[build_stage_mapping(memory_stage)],
    )

    assert len(result) == 1
    payload_raw = result[0]
    if isinstance(payload_raw, (bytes, bytearray)):
        payload_raw = payload_raw.decode("utf-8")
    if isinstance(payload_raw, str):
        payload = json.loads(payload_raw)
    else:
        payload = payload_raw
    assert isinstance(payload, dict)
    pages = [
        {"index": p["index"], "content": "<PAGE_CONTENT>"}
        for p in (payload.get("pages") or [])
    ]
    page_count = len(pages)
    normalized = {
        "pages": pages,
        "metadata": {"pageCount": page_count},
        "errorInformation": payload.get("errorInformation"),
    }

    expected = {
        "pages": pages,
        "metadata": {"pageCount": page_count},
        "errorInformation": None,
    }

    actual_str = json.dumps(normalized, ensure_ascii=False, sort_keys=True)
    expected_str = json.dumps(expected, ensure_ascii=False, sort_keys=True)
    assert actual_str == expected_str


def test_parse_document_docx_round_trip(running_server, memory_stage):
    _write_multipage_docx(memory_stage, "multi.docx")
    client = UDFClient(host="127.0.0.1", port=running_server)
    result = client.call_function(
        "ai_parse_document",
        "multi.docx",
        2048,
        stage_locations=[build_stage_mapping(memory_stage)],
    )

    assert len(result) == 1
    payload_raw = result[0]
    if isinstance(payload_raw, (bytes, bytearray)):
        payload_raw = payload_raw.decode("utf-8")
    if isinstance(payload_raw, str):
        payload = json.loads(payload_raw)
    else:
        payload = payload_raw
    assert isinstance(payload, dict)
    pages = [
        {"index": p["index"], "content": "<PAGE_CONTENT>"}
        for p in (payload.get("pages") or [])
    ]
    page_count = len(pages)
    normalized = {
        "pages": pages,
        "metadata": {"pageCount": page_count},
        "errorInformation": payload.get("errorInformation"),
    }

    expected = {
        "pages": pages,
        "metadata": {"pageCount": page_count},
        "errorInformation": None,
    }

    actual_str = json.dumps(normalized, ensure_ascii=False, sort_keys=True)
    expected_str = json.dumps(expected, ensure_ascii=False, sort_keys=True)
    assert actual_str == expected_str
